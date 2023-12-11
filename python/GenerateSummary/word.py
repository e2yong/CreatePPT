from docx import Document
#python-docx가 설치되어 있어야함
from docx.oxml.ns import qn
import setting

RESULT_DIR = setting.RESULT_DIR

#워드 작성 함수
def Write_word(script_data, filename):
    #워드 파일에 넣은 값과 이름을 받아 진행하는 함수
    #파일 경로를 리턴값으로 지정하여 POST시 파일의 경로를 보내는데 사용함
    doc = Document()
    wordfilename = filename.replace("_extract.txt", "_script.docx")
    word_filepath = RESULT_DIR + "/" + wordfilename
    doc.add_heading('발표대본',level=0)

    #제목은 일단 발표대본으로 임시 지정
    para = doc.add_paragraph()
    run = para.add_run(script_data)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    doc.save(word_filepath)
    return word_filepath